"""Word document parser using python-docx."""

from __future__ import annotations

import io
import time
from typing import Any

import structlog

from app.docprocessor.base import ExtractionResult, ExtractedSection, ExtractedTable

logger = structlog.stdlib.get_logger()


class WordParser:
    """Extract structured data from Word documents using python-docx."""

    MAX_FILE_SIZE = 100_000_000  # 100 MB
    MAX_TEXT_SIZE = 50_000_000  # 50 MB extracted text cap (decompression bomb guard)

    async def parse(
        self,
        file_path: str | None = None,
        file_bytes: bytes | None = None,
        filename: str = "",
    ) -> ExtractionResult:
        """Parse Word document and extract structured content."""
        start = time.monotonic()
        result = ExtractionResult(source_type="word", source_name=filename or file_path or "unknown.docx")

        if file_bytes and len(file_bytes) > self.MAX_FILE_SIZE:
            result.metadata["error"] = f"File too large ({len(file_bytes)} bytes, max {self.MAX_FILE_SIZE})"
            result.extraction_duration_ms = int((time.monotonic() - start) * 1000)
            return result

        try:
            import docx
        except ImportError:
            logger.error("python_docx_not_installed")
            result.metadata["error"] = "python-docx is not installed"
            result.extraction_duration_ms = int((time.monotonic() - start) * 1000)
            return result

        try:
            if file_path:
                doc = docx.Document(file_path)
            else:
                doc = docx.Document(io.BytesIO(file_bytes))  # type: ignore[arg-type]

            # Extract sections from headings and paragraphs
            sections = self._extract_sections(doc)
            result.sections = sections

            # Extract tables
            tables = self._extract_tables(doc)
            result.tables = tables

            # Build raw text (cap to prevent decompression-bomb OOM)
            all_text_parts: list[str] = []
            total_text_len = 0
            for section in sections:
                if section.title:
                    all_text_parts.append(section.title)
                    total_text_len += len(section.title)
                all_text_parts.append(section.content)
                total_text_len += len(section.content)
                if total_text_len > self.MAX_TEXT_SIZE:
                    logger.warning("word_text_size_limit", filename=filename, size=total_text_len)
                    break
            result.raw_text = "\n\n".join(all_text_parts)

            # Metadata
            core = doc.core_properties
            if core.title:
                result.metadata["title"] = core.title
            if core.author:
                result.metadata["author"] = core.author

        except Exception as exc:
            logger.error("word_extraction_failed", error=str(exc), filename=filename)
            result.metadata["error"] = str(exc)

        result.extraction_duration_ms = int((time.monotonic() - start) * 1000)
        return result

    def _extract_sections(self, doc: Any) -> list[ExtractedSection]:
        """Extract sections based on heading hierarchy."""
        sections: list[ExtractedSection] = []
        current_title = ""
        current_level = 1
        current_content: list[str] = []

        _HEADING_LEVELS = {
            "Heading 1": 1,
            "Heading 2": 2,
            "Heading 3": 3,
            "Heading 4": 4,
            "Heading 5": 5,
            "Heading 6": 6,
        }

        _MAX_PARAGRAPHS = 5000
        for para in doc.paragraphs[:_MAX_PARAGRAPHS]:
            style_name = para.style.name if para.style else ""

            if style_name in _HEADING_LEVELS:
                # Save previous section
                if current_content or current_title:
                    sections.append(
                        ExtractedSection(
                            title=current_title,
                            content="\n".join(current_content).strip(),
                            level=current_level,
                        )
                    )
                current_title = para.text.strip()
                current_level = _HEADING_LEVELS[style_name]
                current_content = []
            elif style_name.startswith("List"):
                # Preserve list items with bullet markers
                text = para.text.strip()
                if text:
                    current_content.append(f"- {text}")
            else:
                text = para.text.strip()
                if text:
                    current_content.append(text)

        # Save last section
        if current_content or current_title:
            sections.append(
                ExtractedSection(
                    title=current_title,
                    content="\n".join(current_content).strip(),
                    level=current_level,
                )
            )

        return sections

    def _extract_tables(self, doc: Any) -> list[ExtractedTable]:
        """Extract tables from Word document."""
        extracted: list[ExtractedTable] = []

        for table in doc.tables:
            try:
                rows_data: list[list[str]] = []
                for row in table.rows:
                    row_values = [cell.text.strip() for cell in row.cells]
                    rows_data.append(row_values)

                if not rows_data:
                    continue

                headers = rows_data[0]
                data_rows = rows_data[1:] if len(rows_data) > 1 else []

                extracted.append(
                    ExtractedTable(
                        headers=headers,
                        rows=data_rows,
                    )
                )
            except Exception as exc:
                logger.warning("word_table_extraction_failed", error=str(exc))

        return extracted
